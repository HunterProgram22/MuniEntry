from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def getText(filename):
    "There are formatting issues that need to be fixed."
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)



"""
    def createEntry(self):
        context = self.getDialogFields()
        doc = DocxTemplate("Templates/JuryInstructionsMaster.docx")
        doc.render(context)
        for para in doc.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.save("Saved/Jury_Instructions_Test.docx")
        #Need to us os to get system Path
        os.startfile(PATH + "Saved/Jury_Instructions_Test.docx")
"""
